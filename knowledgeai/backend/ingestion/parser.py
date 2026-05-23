"""
parser.py — Extract raw text from PDF, DOCX, and plain-text files.
             Also extract embedded screenshots/images page by page.
"""
import os
import logging
from typing import Tuple, List, Dict, Optional

logger = logging.getLogger(__name__)


def parse_file(file_path: str) -> Tuple[str, dict]:
    """
    Parse a document and return (text, meta).
    meta keys: file_name, file_type, page_count, page_word_offsets
      page_word_offsets: list of (page_num, start_word_index) for page-aware chunk tagging
    """
    ext = os.path.splitext(file_path)[1].lower()
    file_name = os.path.basename(file_path)

    if ext == ".pdf":
        return _parse_pdf(file_path, file_name)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_path, file_name)
    elif ext in (".txt", ".md", ".rst"):
        return _parse_text(file_path, file_name)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _parse_pdf(path: str, name: str) -> Tuple[str, dict]:
    import pdfplumber
    pages_text = []
    page_word_offsets = []   # (page_num, cumulative_word_start)
    cumulative_words = 0

    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                text = (page.extract_text() or "").strip()
                page_word_offsets.append((i + 1, cumulative_words))
                cumulative_words += len(text.split())
                pages_text.append(text)
    except Exception as e:
        logger.error(f"PDF parse error {path}: {e}")
        raise

    full_text = "\n\n".join(p for p in pages_text if p)
    return full_text, {
        "file_name": name,
        "file_type": "pdf",
        "page_count": page_count,
        "page_word_offsets": page_word_offsets,
    }


def _parse_docx(path: str, name: str) -> Tuple[str, dict]:
    from docx import Document
    try:
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        page_count = max(1, len(paragraphs) // 30)   # rough estimate
    except Exception as e:
        logger.error(f"DOCX parse error {path}: {e}")
        raise

    return full_text, {
        "file_name": name,
        "file_type": "docx",
        "page_count": page_count,
        "page_word_offsets": None,   # DOCX has no reliable page layout
    }


def _parse_text(path: str, name: str) -> Tuple[str, dict]:
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
    except Exception as e:
        logger.error(f"Text parse error {path}: {e}")
        raise

    ext = os.path.splitext(name)[1].lower().lstrip(".")
    return text, {
        "file_name": name,
        "file_type": ext or "txt",
        "page_count": 1,
        "page_word_offsets": None,
    }


# ── Image extraction ────────────────────────────────────────────────────────

def extract_images(file_path: str, doc_id: str, images_dir: str) -> List[Dict]:
    """
    Extract embedded screenshots/images from a document.
    Returns list of dicts: {page_num, filename, filepath}
      page_num is 1-based for PDFs; None for DOCX (no page layout available).
    """
    ext = os.path.splitext(file_path)[1].lower()
    os.makedirs(images_dir, exist_ok=True)

    if ext == ".pdf":
        return _extract_pdf_images(file_path, doc_id, images_dir)
    elif ext in (".docx", ".doc"):
        return _extract_docx_images(file_path, doc_id, images_dir)
    return []


def _extract_pdf_images(path: str, doc_id: str, images_dir: str) -> List[Dict]:
    """Render pages that contain embedded images and save as PNG."""
    import pdfplumber
    results = []
    MAX_PAGES = 60   # safety cap

    try:
        with pdfplumber.open(path) as pdf:
            pages_done = 0
            for i, page in enumerate(pdf.pages):
                if pages_done >= MAX_PAGES:
                    break
                if not page.images:
                    continue   # text-only page — skip
                try:
                    img_filename = f"{doc_id}_p{i + 1}.png"
                    img_path = os.path.join(images_dir, img_filename)
                    if not os.path.exists(img_path):   # don't re-render unchanged docs
                        page.to_image(resolution=120).save(img_path, format="PNG")
                    results.append({
                        "page_num":  i + 1,
                        "filename":  img_filename,
                        "filepath":  img_path,
                    })
                    pages_done += 1
                except Exception as e:
                    logger.warning("Could not render PDF page %d of %s: %s", i + 1, path, e)
    except Exception as e:
        logger.warning("PDF image extraction failed for %s: %s", path, e)

    return results


def _extract_docx_images(path: str, doc_id: str, images_dir: str) -> List[Dict]:
    """Extract inline images from a DOCX file."""
    from docx import Document
    results = []

    try:
        doc = Document(path)
        img_idx = 0
        for rel in doc.part.rels.values():
            if "image" not in rel.reltype:
                continue
            try:
                blob = rel.target_part.blob
                content_type = rel.target_part.content_type   # e.g. "image/png"
                raw_ext = content_type.split("/")[-1].lower()
                ext_map = {"jpeg": "jpg", "jpg": "jpg", "png": "png",
                           "gif": "gif", "bmp": "bmp", "tiff": "tif",
                           "x-emf": "emf", "x-wmf": "wmf"}
                file_ext = ext_map.get(raw_ext, "png")

                # Skip tiny images (icons < 2 KB are rarely screenshots)
                if len(blob) < 2048:
                    continue

                img_filename = f"{doc_id}_img{img_idx}.{file_ext}"
                img_path = os.path.join(images_dir, img_filename)
                if not os.path.exists(img_path):
                    with open(img_path, "wb") as f:
                        f.write(blob)

                results.append({
                    "page_num":  None,   # DOCX: no reliable page number
                    "filename":  img_filename,
                    "filepath":  img_path,
                })
                img_idx += 1
            except Exception as e:
                logger.warning("Could not extract DOCX image %d from %s: %s", img_idx, path, e)

    except Exception as e:
        logger.warning("DOCX image extraction failed for %s: %s", path, e)

    return results
